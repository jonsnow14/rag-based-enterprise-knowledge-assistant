"""Parse PDF / DOCX / XLSX into plain text + lightweight section candidates."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from src.ingestion.discover import SourceFile, slugify
from src.observability import get_logger

log = get_logger(__name__)


@dataclass
class ParsedSection:
    title: str
    text: str
    kind: str = "section"  # section | sheet | table


@dataclass
class ParsedDocument:
    source: SourceFile
    title: str
    doc_id: str
    full_text: str
    sections: List[ParsedSection] = field(default_factory=list)
    effective_date: Optional[str] = None
    version: str = "1.0"
    is_current: bool = True


_HEADING_RE = re.compile(
    r"^(?P<title>(?:\d+(?:\.\d+)*\.?|[A-Z][A-Z0-9 /&-]{2,80}|§\s*\d+[^\n]*))\s*$",
    re.MULTILINE,
)
_DATE_RE = re.compile(
    r"Effective:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4}|\d{4}-\d{2}-\d{2}|January\s+\d{1,2},\s+\d{4})",
    re.I,
)
_VERSION_RE = re.compile(r"Version\s*[: ]\s*([0-9]+(?:\.[0-9]+)*)", re.I)


def _meta_flags(filename: str, text: str) -> tuple[bool, Optional[str], str]:
    """Return is_current, effective_date, version."""
    is_current = True
    lower = filename.lower()
    if "2025" in lower and "pricing" in lower:
        is_current = False
    if "2025" in lower and "2026" not in lower:
        # older rate card pattern
        if "pricing" in lower:
            is_current = False

    effective = None
    m = _DATE_RE.search(text[:2000])
    if m:
        effective = m.group(1).strip()
    elif "2026" in filename:
        effective = "2026-01-01"
    elif "2025" in filename:
        effective = "2025-01-01"

    version = "1.0"
    vm = _VERSION_RE.search(text[:2000])
    if vm:
        version = vm.group(1)
    return is_current, effective, version


def _is_section_heading(stripped: str) -> bool:
    """True for policy headings like '2.1 Annual…', not table cells like '15 days'."""
    if len(stripped) >= 120 or len(stripped) < 5:
        return False
    # Require multi-level or trailing-dot section numbers: "1." "2.1" "3. Subscription"
    # Reject plain "30 days" / "15 days" table cells.
    if re.match(r"^\d+\.\d+", stripped):  # 2.1 …
        return bool(re.search(r"[A-Za-z]", stripped))
    if re.match(r"^\d+\.\s+[A-Za-z]", stripped):  # 1. Purpose
        return True
    if re.match(r"^\d+\s+[A-Z][a-zA-Z]", stripped) and not re.search(
        r"\b(days?|weeks?|months?|years?|hours?|seats?)\b", stripped, re.I
    ):
        return True
    return False


def _split_sections(text: str) -> List[ParsedSection]:
    """Best-effort split on numbered policy headings; keep table rows in-section."""
    lines = text.splitlines()
    sections: List[ParsedSection] = []
    current_title = "Document"
    buf: List[str] = []

    def flush() -> None:
        body = "\n".join(buf).strip()
        if body:
            sections.append(ParsedSection(title=current_title, text=body, kind="section"))

    for line in lines:
        stripped = line.strip()
        if not stripped:
            buf.append(line)
            continue
        if _is_section_heading(stripped):
            flush()
            current_title = stripped[:120]
            buf = []
            continue
        buf.append(line)
    flush()

    if not sections and text.strip():
        sections = [ParsedSection(title="Document", text=text.strip(), kind="section")]
    return sections


def parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n".join(parts)


def parse_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts: List[str] = []
    for p in document.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # Tables are often where limits/requirements live (password length, etc.)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_xlsx(path: Path) -> List[ParsedSection]:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    sections: List[ParsedSection] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: List[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).strip() for c in row]
            if any(cells):
                rows.append("\t".join(cells))
        text = "\n".join(rows).strip()
        if text:
            sections.append(ParsedSection(title=sheet_name, text=text, kind="sheet"))
    wb.close()
    return sections


def parse_file(source: SourceFile) -> ParsedDocument:
    path = source.path
    suffix = path.suffix.lower()
    title = path.stem
    doc_id = f"{source.department.lower()}_{slugify(path.stem)}"

    if suffix == ".pdf":
        full = parse_pdf(path)
        sections = _split_sections(full)
    elif suffix == ".docx":
        full = parse_docx(path)
        sections = _split_sections(full)
    elif suffix == ".xlsx":
        sections = parse_xlsx(path)
        full = "\n\n".join(f"## {s.title}\n{s.text}" for s in sections)
    else:
        raise ValueError(f"unsupported type: {suffix}")

    # strip repeated footer noise lightly
    full = re.sub(r"Northwind Traders, Inc\.\s*—\s*Internal Use Only\s*Page\s*\d+", " ", full)
    full = re.sub(r"[ \t]+\n", "\n", full)

    is_current, effective, version = _meta_flags(source.filename, full)
    if not sections:
        sections = [ParsedSection(title="Document", text=full or title, kind="section")]

    log.info(
        "parsed %s dept=%s sections=%s chars=%s current=%s",
        source.filename,
        source.department,
        len(sections),
        len(full),
        is_current,
    )
    return ParsedDocument(
        source=source,
        title=title,
        doc_id=doc_id,
        full_text=full,
        sections=sections,
        effective_date=effective,
        version=version,
        is_current=is_current,
    )
